from mailmerge import MailMerge
from datetime import date
from docxtpl import DocxTemplate,InlineImage
from docx.shared import *
from docx import *
def getquestions():
    # -- coding: utf-8 --
    """
    Created on Thu Jan 23 16:02:36 2020

    @author: Jai Kumawat
    """
    import datetime
    d={}

    lst=[]
    print("enter occurance of offence :")
    a=input("Enter day  : ")
    d["day1"]=a
    a=input("Enter date (eg: dd/mm/yyyy) : ")
    d["year1"]=a
    a=input("Enter time (eg: in hours) : ")
    d["hours1"]=a
    print("\n Information received on :")

    x=datetime.datetime.now().strftime("%x")
    print(x)
    d["year2"]= x
    x=datetime.datetime.now().strftime("%X")
    d["hours2"]=x


    print("\n Enter occurance of offence")
    a=input("Direction and distance from police station : ")
    d["direction_distance"]=a
    a=input("address : ")
    d["address"]=a
    print("\n Complainant/informant details :")
    a=input("name : ")
    d["name2"]=a
    a=input("Father/Husband's name : ")
    d["name3"]=a
    a=input("Nationality : ")
    d["nationality"]=a
    a=input("Address : ")
    d["address1"]=a
    a=input("Details of known/unknown/ suspect/ accused with full particulars : ")
    d["details"]=a
    a=input("Reasons for delay in reporting the crime : ")
    d["reason"]=a
    b=input("Was any property stone/involved ? (yes or no) : ")
    if(b=="yes"or b=="Yes" or b=="YES"):
        a=input("Particulars of property stolen/involved : ")
        d["reason1"]=a
        a=input("Total value of property stolen/involved : ")
        d["value"]=a
    else:
        a=input("Inquest report/U.D : Case No., if any : ")
        d["fill"]=a
        a=input("FIR contents : ")
        d["fill1"]=a
    return d


def generate():
    info=getquestions()
    print(info)
    template="./template.docx"
    doc=MailMerge(template)
    doc.merge_pages([info])
    doc.write("./output.docx")
    document = Document('./output.docx')
    document.add_picture("D:\SIH2020\output..jpg",width=Inches(1),height=Inches(0.5))
    document.add_paragraph('Signature')
    document.save("./output.docx")


generate()